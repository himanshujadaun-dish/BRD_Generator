import streamlit as st
import openai
from docx import Document
from datetime import date
import io

# ------------------------------------------------------------
# SETUP
# ------------------------------------------------------------
st.set_page_config(page_title="BRD Generator", layout="wide")
openai.api_key = st.secrets["openai_api_key"]


# ------------------------------------------------------------
# GPT BRD GENERATION FUNCTION
# ------------------------------------------------------------
def generate_brd_text(data):
    prompt = f"""
You are an expert Business Systems Analyst at DISH Wireless / Boost Mobile.

Generate a Business Requirements Document (BRD) for a Tableau dashboard 
using the EXACT structure and section names below:

------------------------------------------
Business Requirements Document (Dashboard Request)

Project / Dashboard Name:
Date Created:
Requested By (Business Team):
Prepared By (Analyst):
Version:

1️⃣ Business Overview
• Business Problem / Need
• Business Goal / Outcome
• Scope (In-Scope)
• Out of Scope
• Expected Frequency

2️⃣ Key Stakeholders
(Table with: Role, Name, Department / Notes)

3️⃣ Data Inputs
(Table with: Source System/Table, Description, Frequency, Owner)

4️⃣ Dashboard Requirements
(Table with: Dashboard Section/Visualization, Description/Purpose,
Key Metrics or Fields, Filters Required, Drilldown Needed?)

5️⃣ Business Rules / Calculations
(Table with: Metric, Definition / Formula, Notes)

6️⃣ Expected Outputs
(Table with: Deliverable, Format/Platform, Frequency, Audience)

7️⃣ Validation & Sign-off
(Table with: Step, Responsible, Criteria, Status)

8️⃣ Notes / Attachments

9️⃣ Control Data & Validation Sources
(Table with: Control Report/Source, Description / Purpose,
Business Owner, Validation Method, Frequency)
------------------------------------------

Use this DATA to populate each section:

{data}

Return only structured text that matches the template.
"""

    response = openai.chat.completions.create(
        model="gpt-4o",
        messages=[{"role": "user", "content": prompt}]
    )

    return response.choices[0].message.content


# ------------------------------------------------------------
# FUNCTION: Create Word Document With Correct Tables
# ------------------------------------------------------------
def create_brd_docx(brd_text, form_inputs):
    doc = Document()

    # Header
    doc.add_heading("Business Requirements Document (Dashboard Request)", level=1)

    doc.add_paragraph(f"Project / Dashboard Name: {form_inputs['project_name']}")
    doc.add_paragraph(f"Date Created: {form_inputs['date_created']}")
    doc.add_paragraph(f"Requested By (Business Team): {form_inputs['requested_by']}")
    doc.add_paragraph(f"Prepared By (Analyst): {form_inputs['prepared_by']}")
    doc.add_paragraph(f"Version: {form_inputs['version']}")
    doc.add_paragraph("")

    # 1️⃣ Business Overview
    doc.add_heading("1️⃣ Business Overview", level=2)
    doc.add_paragraph("Business Problem / Need:")
    doc.add_paragraph(form_inputs["business_problem"])

    doc.add_paragraph("\nBusiness Goal / Outcome:")
    doc.add_paragraph(form_inputs["business_goal"])

    doc.add_paragraph("\nScope (In-Scope):")
    doc.add_paragraph(form_inputs["in_scope"])

    doc.add_paragraph("\nOut of Scope:")
    doc.add_paragraph(form_inputs["out_of_scope"])

    doc.add_paragraph("\nExpected Frequency:")
    doc.add_paragraph(form_inputs["frequency"])

    # 2️⃣ Key Stakeholders
    doc.add_heading("2️⃣ Key Stakeholders", level=2)
    table = doc.add_table(rows=1, cols=3)
    hdr = table.rows[0].cells
    hdr[0].text = "Role"
    hdr[1].text = "Name"
    hdr[2].text = "Department / Notes"

    for row in form_inputs["stakeholders"]:
        cells = table.add_row().cells
        cells[0].text = row["role"]
        cells[1].text = row["name"]
        cells[2].text = row["dept"]

    # 3️⃣ Data Inputs
    doc.add_heading("3️⃣ Data Inputs", level=2)
    table = doc.add_table(rows=1, cols=4)
    table.rows[0].cells[0].text = "Source System/Table"
    table.rows[0].cells[1].text = "Description"
    table.rows[0].cells[2].text = "Frequency"
    table.rows[0].cells[3].text = "Owner"

    for row in form_inputs["data_inputs"]:
        cells = table.add_row().cells
        cells[0].text = row["source"]
        cells[1].text = row["description"]
        cells[2].text = row["frequency"]
        cells[3].text = row["owner"]

    # 4️⃣ Dashboard Requirements
    doc.add_heading("4️⃣ Dashboard Requirements", level=2)
    table = doc.add_table(rows=1, cols=5)
    table.rows[0].cells[0].text = "Dashboard Section"
    table.rows[0].cells[1].text = "Description / Purpose"
    table.rows[0].cells[2].text = "Key Metrics / Fields"
    table.rows[0].cells[3].text = "Filters Required"
    table.rows[0].cells[4].text = "Drilldown Needed?"

    for row in form_inputs["dash_reqs"]:
        cells = table.add_row().cells
        cells[0].text = row["section"]
        cells[1].text = row["purpose"]
        cells[2].text = row["metrics"]
        cells[3].text = row["filters"]
        cells[4].text = row["drill"]

    # 5️⃣ Business Rules
    doc.add_heading("5️⃣ Business Rules / Calculations", level=2)
    table = doc.add_table(rows=1, cols=3)
    table.rows[0].cells[0].text = "Metric"
    table.rows[0].cells[1].text = "Definition / Formula"
    table.rows[0].cells[2].text = "Notes"

    for row in form_inputs["business_rules"]:
        cells = table.add_row().cells
        cells[0].text = row["metric"]
        cells[1].text = row["formula"]
        cells[2].text = row["notes"]

    # 6️⃣ Expected Outputs
    doc.add_heading("6️⃣ Expected Outputs", level=2)
    table = doc.add_table(rows=1, cols=4)
    table.rows[0].cells[0].text = "Deliverable"
    table.rows[0].cells[1].text = "Format / Platform"
    table.rows[0].cells[2].text = "Frequency"
    table.rows[0].cells[3].text = "Audience"

    for row in form_inputs["expected_outputs"]:
        cells = table.add_row().cells
        cells[0].text = row["deliverable"]
        cells[1].text = row["format"]
        cells[2].text = row["freq"]
        cells[3].text = row["audience"]

    # 7️⃣ Validation & Sign-off
    doc.add_heading("7️⃣ Validation & Sign-off", level=2)
    table = doc.add_table(rows=1, cols=4)
    table.rows[0].cells[0].text = "Step"
    table.rows[0].cells[1].text = "Responsible"
    table.rows[0].cells[2].text = "Criteria"
    table.rows[0].cells[3].text = "Status"

    for row in form_inputs["validation"]:
        cells = table.add_row().cells
        cells[0].text = row["step"]
        cells[1].text = row["owner"]
        cells[2].text = row["criteria"]
        cells[3].text = row["status"]

    # 8️⃣ Notes / Attachments
    doc.add_heading("8️⃣ Notes / Attachments", level=2)
    doc.add_paragraph(form_inputs["notes"])

    # 9️⃣ Control Data
    doc.add_heading("9️⃣ Control Data & Validation Sources", level=2)
    table = doc.add_table(rows=1, cols=5)
    table.rows[0].cells[0].text = "Control Report / Source"
    table.rows[0].cells[1].text = "Description / Purpose"
    table.rows[0].cells[2].text = "Business Owner"
    table.rows[0].cells[3].text = "Validation Method"
    table.rows[0].cells[4].text = "Frequency"

    for row in form_inputs["control_data"]:
        cells = table.add_row().cells
        cells[0].text = row["source"]
        cells[1].text = row["description"]
        cells[2].text = row["owner"]
        cells[3].text = row["method"]
        cells[4].text = row["frequency"]

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ------------------------------------------------------------
# STREAMLIT UI
# ------------------------------------------------------------
st.title("📊 Boost Mobile BRD Generator")
st.markdown("Generate a BRD identical to your team's official template.")

# BASIC INFO
project_name = st.text_input("Project Name")
requested_by = st.text_input("Requested By (Business Team)")
prepared_by = st.text_input("Prepared By (Analyst)")
version = st.text_input("Version", value="1.0")
date_created = st.date_input("Date Created", value=date.today())

# BUSINESS OVERVIEW
st.subheader("1️⃣ Business Overview")
business_problem = st.text_area("Business Problem")
business_goal = st.text_area("Business Goal")
in_scope = st.text_area("In Scope")
out_of_scope = st.text_area("Out of Scope")
frequency = st.selectbox("Expected Frequency", ["Daily", "Weekly", "Monthly"])

# MULTI-ROW TABLE INPUTS
def table_input(label, cols):
    st.write(f"### {label}")
    rows = st.text_area(f"Enter rows in CSV format ({', '.join(cols)}):")
    parsed_rows = []

    if rows.strip():
        for line in rows.split("\n"):
            parts = [p.strip() for p in line.split(",")]
            if len(parts) == len(cols):
                parsed_rows.append(dict(zip(cols, parts)))

    return parsed_rows


stakeholders = table_input("2️⃣ Key Stakeholders", ["role", "name", "dept"])
data_inputs = table_input("3️⃣ Data Inputs", ["source", "description", "frequency", "owner"])
dash_reqs = table_input("4️⃣ Dashboard Requirements", ["section", "purpose", "metrics", "filters", "drill"])
business_rules = table_input("5️⃣ Business Rules", ["metric", "formula", "notes"])
expected_outputs = table_input("6️⃣ Expected Outputs", ["deliverable", "format", "freq", "audience"])
validation = table_input("7️⃣ Validation Steps", ["step", "owner", "criteria", "status"])
control_data = table_input("9️⃣ Control Data", ["source", "description", "owner", "method", "frequency"])

notes = st.text_area("8️⃣ Notes / Attachments")

# GENERATE BRD
if st.button("Generate BRD (.docx)"):
    with st.spinner("Generating BRD..."):
        form_inputs = {
            "project_name": project_name,
            "requested_by": requested_by,
            "prepared_by": prepared_by,
            "version": version,
            "date_created": date_created,
            "business_problem": business_problem,
            "business_goal": business_goal,
            "in_scope": in_scope,
            "out_of_scope": out_of_scope,
            "frequency": frequency,
            "stakeholders": stakeholders,
            "data_inputs": data_inputs,
            "dash_reqs": dash_reqs,
            "business_rules": business_rules,
            "expected_outputs": expected_outputs,
            "validation": validation,
            "notes": notes,
            "control_data": control_data,
        }

        buffer = create_brd_docx("", form_inputs)

        st.success("BRD Created Successfully!")
        st.download_button(
            label="📥 Download BRD Document",
            data=buffer,
            file_name="BRD.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
